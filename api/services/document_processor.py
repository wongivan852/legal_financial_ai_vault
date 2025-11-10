"""
Document Processing Service
Handles text extraction from PDF and DOCX files
"""

import logging
from pathlib import Path
from typing import List, Optional, Union
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from models.document import Document
from security.encryption import encryption_service

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing and extracting text from documents"""

    @staticmethod
    def extract_text_from_pdf(file_path: Union[str, Path]) -> str:
        """
        Extract text from PDF file

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        text_content = []

        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            if text_content:
                return "\n\n".join(text_content)

        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}, trying PyPDF2: {e}")

        try:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            return "\n\n".join(text_content)

        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise

    @staticmethod
    def extract_text_from_docx(file_path: Union[str, Path]) -> str:
        """
        Extract text from DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            file_path = Path(file_path)
            doc = DocxDocument(file_path)

            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_content.append(" | ".join(row_text))

            return "\n\n".join(text_content)

        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise

    @staticmethod
    def extract_text_from_txt(file_path: Union[str, Path]) -> str:
        """
        Extract text from TXT file

        Args:
            file_path: Path to TXT file

        Returns:
            File content as string
        """
        try:
            file_path = Path(file_path)
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            raise

    @staticmethod
    def extract_text(file_path: Union[str, Path], file_type: str) -> str:
        """
        Extract text from document based on file type

        Args:
            file_path: Path to file
            file_type: File extension (.pdf, .docx, .txt)

        Returns:
            Extracted text content
        """
        file_type = file_type.lower()

        if file_type == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_type in ['.docx', '.doc']:
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif file_type == '.txt':
            return DocumentProcessor.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def chunk_document(
        text: str,
        chunk_size: int = 20000,
        overlap: int = 500
    ) -> List[str]:
        """
        Split document into overlapping chunks

        Args:
            text: Document text to chunk
            chunk_size: Maximum characters per chunk
            overlap: Number of overlapping characters between chunks

        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]

            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)

                if break_point > chunk_size * 0.7:  # If break point is reasonable
                    chunk = chunk[:break_point + 1]
                    end = start + break_point + 1

            chunks.append(chunk.strip())
            start = end - overlap

        return chunks

    @staticmethod
    async def get_document_text(document_id: str, db: Session) -> str:
        """
        Get extracted text for a document

        Args:
            document_id: Document ID
            db: Database session

        Returns:
            Document text content
        """
        document = db.query(Document).filter(Document.id == document_id).first()

        if not document:
            raise ValueError(f"Document not found: {document_id}")

        # If text already extracted, return it
        if document.text_content:
            return document.text_content

        # Otherwise, extract text from file
        if document.encrypted:
            # Decrypt file to temporary location
            decrypted_data = encryption_service.read_encrypted_file(document.storage_path)
            # For simplicity, assume we can process from bytes
            # In production, you might want to write to temp file
            raise NotImplementedError("Extracting from encrypted files needs temp file handling")
        else:
            text = DocumentProcessor.extract_text(document.storage_path, document.file_type)
            return text

    @staticmethod
    def count_words(text: str) -> int:
        """Count words in text"""
        return len(text.split())

    @staticmethod
    def count_pages_pdf(file_path: Union[str, Path]) -> int:
        """
        Count pages in PDF file

        Args:
            file_path: Path to PDF file

        Returns:
            Number of pages
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        except Exception as e:
            logger.error(f"Error counting PDF pages: {e}")
            return 0

    @staticmethod
    def get_document_metadata(file_path: Union[str, Path], file_type: str) -> dict:
        """
        Extract metadata from document

        Args:
            file_path: Path to file
            file_type: File extension

        Returns:
            Dictionary with metadata
        """
        metadata = {}

        try:
            if file_type == '.pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['page_count'] = len(pdf_reader.pages)

                    if pdf_reader.metadata:
                        metadata['title'] = pdf_reader.metadata.get('/Title')
                        metadata['author'] = pdf_reader.metadata.get('/Author')
                        metadata['subject'] = pdf_reader.metadata.get('/Subject')
                        metadata['creator'] = pdf_reader.metadata.get('/Creator')

        except Exception as e:
            logger.error(f"Error extracting metadata: {e}")

        return metadata
